"""Part 1: Helper functions for Word doc generation."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(BASE, "docs", "figures")

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for i in range(1,4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0, 51, 102)
        hs.font.bold = True
    return doc

def add_title_page(doc):
    for _ in range(6): doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FOOD DELIVERY DELAY &\nCUSTOMER RATING ANALYSIS")
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor(0,51,102)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Complete Project Documentation & Panel Presentation Guide")
    r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(100,100,100)
    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = d.add_run("Data Science Project | May 2026")
    r3.font.size = Pt(12)
    doc.add_page_break()

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

def add_fig(doc, name, caption):
    path = os.path.join(FIG_DIR, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.font.size = Pt(9); r.font.italic = True
